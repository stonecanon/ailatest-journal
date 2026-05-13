from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "generated" / "未来城市研究院理4六楼办公空间使用申请_简版一页.docx"

FONT_CN = "仿宋"
FONT_HEAD = "黑体"
INK = RGBColor(0x00, 0x00, 0x00)


def set_font(run, name=FONT_CN, size=12, bold=False):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
    run.font.size = Pt(size)
    run.bold = bold
    run.font.color.rgb = INK


def add_run(paragraph, text, name=FONT_CN, size=12, bold=False):
    run = paragraph.add_run(text)
    set_font(run, name=name, size=size, bold=bold)
    return run


def para(doc, text="", first_line=False, after=6, line=1.45, align=None):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(after)
    pf.line_spacing = line
    if first_line:
        pf.first_line_indent = Pt(24)
    if align is not None:
        p.alignment = align
    if text:
        add_run(p, text)
    return p


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_borders(table):
    borders = table._tbl.tblPr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        table._tbl.tblPr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = borders.find(qn(f"w:{edge}"))
        if el is None:
            el = OxmlElement(f"w:{edge}")
            borders.append(el)
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "4")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), "BFBFBF")


def set_cell(cell, text, bold=False, fill=None, align=WD_ALIGN_PARAGRAPH.LEFT):
    cell.text = ""
    if fill:
        shade_cell(cell, fill)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    p = cell.paragraphs[0]
    p.alignment = align
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.2
    add_run(p, text, name=FONT_HEAD if bold else FONT_CN, size=11, bold=bold)


def build():
    doc = Document()
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.2)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.7)
    section.right_margin = Cm(2.7)

    styles = doc.styles
    styles["Normal"].font.name = FONT_CN
    styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_CN)
    styles["Normal"].font.size = Pt(12)

    title = para(doc, align=WD_ALIGN_PARAGRAPH.CENTER, after=18, line=1.1)
    add_run(title, "申请", name=FONT_HEAD, size=22, bold=True)

    p = para(doc, "工程学院：", after=8, line=1.45)
    p.runs[0].font.bold = True
    p.runs[0]._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_HEAD)

    para(
        doc,
        "未来城市研究院原隶属于国土空间规划学院。因日常办公、教学研讨及相关工作需要，现申请使用工程学院理4六楼602、609、610办公室作为教室及办公用房。",
        first_line=True,
        after=6,
    )
    para(
        doc,
        "具体房间及人员安排如下：",
        first_line=True,
        after=5,
    )

    table = doc.add_table(rows=4, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.allow_autofit = False
    set_borders(table)
    headers = ["房间", "用途", "使用人员"]
    for i, text in enumerate(headers):
        set_cell(table.rows[0].cells[i], text, bold=True, fill="F2F2F2", align=WD_ALIGN_PARAGRAPH.CENTER)
    rows = [
        ("602", "办公室", "徐丹华（C18595副教授）、邬佳婧（C20565讲师）"),
        ("609", "办公室", "周淼（C16048副教授）、李沛文（C16204工程师）"),
        ("610", "办公室", "[REDACTED NAME]、赵宇杰（C24510讲师）"),
    ]
    for row, values in zip(table.rows[1:], rows):
        for i, value in enumerate(values):
            set_cell(row.cells[i], value, align=WD_ALIGN_PARAGRAPH.CENTER if i < 2 else WD_ALIGN_PARAGRAPH.LEFT)
        row.cells[0].width = Cm(2.3)
        row.cells[1].width = Cm(2.8)
        row.cells[2].width = Cm(10.0)

    para(
        doc,
        "使用期间，研究院将遵守学校及工程学院有关办公用房、消防安全、资产管理、门禁管理等规定，做好日常维护和安全管理，并配合学院后续空间统筹和管理要求。",
        first_line=True,
        after=6,
    )
    para(
        doc,
        "以上申请，恳请工程学院审核同意，并协助办理相关签字盖章手续。",
        first_line=True,
        after=14,
    )

    para(doc, "特此申请。", first_line=True, after=18)

    para(doc, "申请单位：未来城市研究院（盖章）", align=WD_ALIGN_PARAGRAPH.RIGHT, after=10)
    para(doc, "日期：        年      月      日", align=WD_ALIGN_PARAGRAPH.RIGHT, after=0)

    doc.core_properties.title = "未来城市研究院办公空间使用申请"
    doc.core_properties.author = "Codex"
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
